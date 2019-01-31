from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import time
import math


class Extractor(object):
    _document = None
    _story = None
    _indexes = None
    _page_number = None
    spec_pre = None
    daily_clean = None
    monthly_clean = None
    all_txt_elements = None


    def __init__(self, doc_file):
        """
        :param doc_file: this required to pass to the python-docx class object Document. It should be:
        <DRIVE_LETTER>:\path\to\file.docx
        OR
        /path/to/file.docx

        When creating the extractor object, we need to gather some of the elements so that
        processing them afterwards becomes quicker and seamless.
        I decided to use another class object to create a duplicate of python-docx's Document
        so that the original contents of the object are not tampered with at all.
        """
        self._document = Document(doc_file)
        self._paragraphs = self._document.paragraphs
        self._indexes = [i for i, v in enumerate(self._document.paragraphs) if v.text.isupper()]
        self.all_txt_elements = self._document.element.xpath('//w:t')

    def _find_headers(self):
        indexes = []
        for i, v in enumerate(self._document.paragraphs):
            if v.text.isupper():
                indexes.append(i)
                print(self._document.paragraphs[i].text)
        self._indexes = indexes

    def _para_print(self):
        for para in self._document.paragraphs:
            print(para.text)

    def set__spec_pre(self):
        first = 0
        last = 0
        discover = 0
        for index, para in enumerate(self._document.paragraphs):
            if "SPECIAL PRECAUTIONS" in para.text and para.text.isupper():
                discover = 1
                first = index+1
            if discover == 1 and para.text.strip() == '' or "DAILY CLEANING METHOD" in para.text:
                last = index-1
                break
        if first == last:
            self.spec_pre = self._document.paragraphs[first]
        elif first != last:
            self.spec_pre = [para.text for para in self._document.paragraphs[first:last]]

    def set__daily_method(self):
        first = 0
        last = 0
        discover = 0
        for index, para in enumerate(self._document.paragraphs):
            if "DAILY CLEANING METHOD" in para.text and para.text.isupper():
                discover = 1
                first = index + 1
            if discover == 1 and para.text.strip() == '' or 'MONTHLY CLEANING METHOD' in para.text \
                    or 'Photo No.' in para.text:
                last = index - 1
                break
        # print(first, last, discover)
        self.daily_clean = [para.text for para in self._document.paragraphs[first:last]]

    def set__story(self):
        self._story = [x for x in extractor.iter_block_items()]

    def iter_block_items(self):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(self._document, _Document):
            parent_elm = self._document.element.body
            # print(parent_elm.xml)
        elif isinstance(self._document, _Cell):
            parent_elm = self._document._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, self._document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, self._document)


def now(time_now=time.time()):
    return time.strftime("%b-%d-%Y_%H:%M:%S", time.gmtime(time_now))


if __name__ == '__main__':
    start = time.time()

    extractor = Extractor('D:\Anaconda3\\feuadal_kingdom\CICPreview. Butchery.pdf.docx')

    count = 0

    extractor.set__spec_pre()

    extractor.set__daily_method()

    extractor.set__story()

    # items = [x for x in extractor.iter_block_items()]

    for i, item in enumerate(extractor._story):
        if hasattr(item, 'text'):
            if len(item.text) > 3:
                if "Area - Description" in item.text:
                    # print(item.text)
                    count += 1
        if hasattr(item, 'rows'):
            for row in item.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if len(para.text) > 3:
                            # print(para.text)
                            if "Area - Description" in para.text:
                                count += 1
    # page_list
    # for index, value in enumerate(extractor.all_txt_elements):
    #     if "Area - Description" in value.text:
    #         count += 1
    # print('Page in line. Total = {}'.format(count))

    finish = time.time()
    elapse = finish - start
    print('{} Total Area - Description matches read in {}'.format(count, elapse))
    # sections = extractor._document.sections
    # for section in sections:
    #     print(section.start_type)
    # docx2txt notes:
    # lines = [item for item in
    # re.split('\n|\t',docx2txt.process('D:\Anaconda3\\feuadal_kingdom\CICPreview. Butchery.pdf.docx')) if len(item) > 1]
    for x in pages_dict:
        # print(lines[pages_dict[x][0]:pages_dict[x][1]])
        if len(lines[pages_dict[x][0]:pages_dict[x][1]]) > 1:
            print(lines[pages_dict[x][0]:pages_dict[x][1]][4])
    for i, y in enumerate(lines[pages_dict[x][0]:pages_dict[x][1]]):
        if "DAILY CLEANING" in y:
            first = i
        if "Photo No." in y:
            last = i
        if first and last:
            print(lines[pages_dict[x][0]:pages_dict[x][1]][first:last])
    for index, line in enumerate(lines):
        if "Area - Description" in line:
            first = index
        if "HYGBU/" in line:
            last = index
        if first != 0 and last != 0:
            pages_dict["page_" + str(index)] = [first, last]
            first = 0
            last = 0
